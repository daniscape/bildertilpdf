def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None



import os
from os import listdir
from os.path import isfile, join
from docx import Document
from docx.shared import Inches


t_w = 2
i = 0


path = os.path.abspath('..') + "\\Photos\\"

### Import the photos to a list, get the length
photos = [f for f in listdir(path) if isfile(join(path, f))]
num_photos = len(photos)


document = Document('template.docx')
delete_paragraph(document.paragraphs[0])
table = document.add_table(1,t_w)
table.style = 'blank'
table.autofit = True

while(len(table.rows) <= num_photos/t_w*2):
    table.add_row()


    p = table.rows[i/t_w*2].cells[i%t_w].paragraphs[0]
    r = p.add_run()
    

    r.add_picture(path + photo, width=Inches(6.0/t_w))
    i += 1


document.save('output.docx')
